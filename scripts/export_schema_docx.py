import sqlite3
import os
from docx import Document

DB_PATH = r"D:\Rpi5.Server 2\instance\Data.db"
OUT_PATH = r"D:\Rpi5.Server 2\docs\Data_schema.docx"

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()
cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name;")
tables = [r[0] for r in cur.fetchall()]

doc = Document()

for t in tables:
    doc.add_heading(f'Table: {t}', level=2)

    # Primary keys
    cur.execute(f"PRAGMA table_info('{t}')")
    cols = cur.fetchall()  # cid,name,type,notnull,dflt_value,pk
    pk = [c[1] for c in cols if c[5] > 0]
    doc.add_paragraph('Primary Key: ' + (', '.join(pk) if pk else '(none)'))

    # Foreign keys
    cur.execute(f"PRAGMA foreign_key_list('{t}')")
    fks = cur.fetchall()
    if fks:
        fk_texts = []
        for fk in fks:
            # fk columns: id, seq, table, from, to, on_update, on_delete, match
            fk_texts.append(f"{fk[2]}({fk[3]}->{fk[4]})")
        doc.add_paragraph('Foreign Keys: ' + ', '.join(fk_texts))
    else:
        doc.add_paragraph('Foreign Keys: (none)')

    # Columns table
    tbl = doc.add_table(rows=1, cols=5)
    hdr = tbl.rows[0].cells
    hdr[0].text = 'Data Item'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Validation'
    hdr[3].text = 'Sample data'
    hdr[4].text = 'Notes'

    # Fetch one sample row if available
    cur.execute(f"SELECT * FROM '{t}' LIMIT 1")
    sample_row = cur.fetchone()

    for i, c in enumerate(cols):
        name = c[1]
        dtype = c[2]
        sample = ''
        if sample_row:
            try:
                val = sample_row[i]
                sample = '' if val is None else str(val)
            except Exception:
                sample = ''
        row_cells = tbl.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = dtype
        row_cells[2].text = ''
        row_cells[3].text = sample
        row_cells[4].text = ''

    doc.add_paragraph('')

conn.close()
doc.save(OUT_PATH)
print(f'Wrote {OUT_PATH}')
